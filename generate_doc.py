from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Color palette ─────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1E, 0x3A, 0x5F)   # #1E3A5F
BLUE_MID   = RGBColor(0x26, 0x5D, 0xA6)   # #265DA6
BLUE_LIGHT = RGBColor(0xE8, 0xF0, 0xFB)   # #E8F0FB
ORANGE     = RGBColor(0xF5, 0x9E, 0x0B)   # #F59E0B
GRAY_TEXT  = RGBColor(0x44, 0x44, 0x44)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top='single', bottom='single', left='single', right='single', size=6, color='265DA6'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   val)
        el.set(qn('w:sz'),    str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_table_border(table):
    tbl    = table._tbl
    tblPr  = tbl.tblPr
    tblBrd = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BBCCE0')
        tblBrd.append(el)
    tblPr.append(tblBrd)

def add_heading_block(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18) if level == 1 else Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(13)
        run.font.color.rgb = BLUE_DARK
    else:
        run.font.size  = Pt(11)
        run.font.color.rgb = BLUE_MID
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.color.rgb = GRAY_TEXT
    return p

def add_service_table(doc, rows, col_headers=('Service', 'Description', 'Utilisateurs')):
    """rows = list of (service, description, users)"""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # header row
    hdr = table.rows[0]
    hdr.height = Cm(0.8)
    widths = [Cm(6.2), Cm(8.2), Cm(4.2)]
    for i, (cell, header, w) in enumerate(zip(hdr.cells, col_headers, widths)):
        cell.width = w
        set_cell_bg(cell, '1E3A5F')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size  = Pt(9.5)
        run.font.color.rgb = WHITE

    # data rows
    for idx, (svc, desc, users) in enumerate(rows):
        row  = table.add_row()
        bg   = 'E8F0FB' if idx % 2 == 0 else 'FFFFFF'
        vals = [svc, desc, users]
        for i, (cell, val) in enumerate(zip(row.cells, vals)):
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY_TEXT
            if i == 0:
                run.bold = True

    set_table_border(table)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(2)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('COMPLEXE SCOLAIRE LE SAGE')
r.bold = True
r.font.size  = Pt(11)
r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('The Wise School International — Kinshasa, RDC')
r.font.size  = Pt(9.5)
r.font.color.rgb = GRAY_TEXT

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SCHOOLSAFE v3.0')
r.bold = True
r.font.size  = Pt(22)
r.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Catalogue des Services & Fonctionnalités')
r.font.size  = Pt(13)
r.font.color.rgb = BLUE_MID

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Édition : {datetime.date.today().strftime("%B %Y")}')
r.font.size  = Pt(9)
r.font.color.rgb = GRAY_TEXT
r.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# INTRO
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, 'À propos de SchoolSafe v3.0', level=1)
add_body(doc,
    'SchoolSafe v3.0 est le système de gestion scolaire intégré du Complexe Scolaire Le Sage / '
    'The Wise School International. Il couvre l\'ensemble des opérations administratives, '
    'pédagogiques, financières et sécuritaires de l\'établissement en temps réel.')
add_body(doc,
    'L\'application est accessible depuis tout navigateur (web, mobile, tablette) et fonctionne '
    'également hors connexion grâce au chiffrement AES-256-GCM local. Elle est bilingue '
    'français / anglais pour les enseignants.')
doc.add_paragraph()

# rôles summary
add_heading_block(doc, 'Rôles utilisateurs', level=2)
roles = [
    ('Direction 1',    'Administrateur général — accès complet à tous les modules'),
    ('Direction 2',    'Responsable pédagogique — académique, approbations, calendrier'),
    ('Direction 3',    'Caissier — finances, recettes, dépenses, rapports'),
    ('Enseignant',     'Professeur — sa classe, notes, devoirs, présences, messagerie'),
    ('Parent',         'Tuteur — suivi enfant, paiements, devoirs, convocations'),
    ('Gardien',        'Sécurité — portail entrée/sortie, scans QR, alertes'),
]
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
hdr = table.rows[0]
for cell, txt in zip(hdr.cells, ['Rôle', 'Périmètre']):
    set_cell_bg(cell, '265DA6')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
for idx, (role, desc) in enumerate(roles):
    row = table.add_row()
    bg  = 'E8F0FB' if idx % 2 == 0 else 'FFFFFF'
    for i, (cell, val) in enumerate(zip(row.cells, [role, desc])):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9); r.font.color.rgb = GRAY_TEXT
        if i == 0: r.bold = True
set_table_border(table)
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. TABLEAUX DE BORD
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '1. Tableaux de bord personnalisés', level=1)
add_body(doc, 'Chaque rôle dispose d\'un tableau de bord adapté à ses responsabilités, '
              'avec des indicateurs clés, des raccourcis et des alertes en temps réel.')
add_service_table(doc, [
    ('Tableau de bord Direction 1', 'Vue d\'ensemble : présences du jour, paiements, scans portail, solde caisse, raccourcis rapides', 'Direction 1'),
    ('Tableau de bord Direction 2', 'Focus pédagogique : approbations en attente, présences par classe, top élèves, absences récentes', 'Direction 2'),
    ('Tableau de bord Caisse',      'Opérations financières du jour : enregistrer paiements, dépenses, consulter le rapport journalier', 'Direction 3'),
    ('Tableau de bord Enseignant',  'Résumé de sa classe : présences, devoirs à venir (7 jours), rattrapages, chat co-titulaire', 'Enseignant'),
    ('Tableau de bord Gardien',     'Accès rapide au scanner entrée/sortie, alertes non autorisées, notifications non lues', 'Gardien'),
    ('Tableau de bord Parent',      'Aperçu enfant : présences, devoirs, emploi du temps, frais scolaires, dernières activités', 'Parent'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 2. GESTION DES ÉLÈVES
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '2. Gestion des élèves', level=1)
add_body(doc, 'Dossier complet pour chaque élève : identité, photo, classe, parents, état de paiement, carte scolaire.')
add_service_table(doc, [
    ('Liste & recherche des élèves',    'Parcourir, filtrer et rechercher tous les élèves (nom, matricule, classe, statut)', 'Direction 1, Direction 2'),
    ('Créer un dossier élève',           'Enregistrer un nouvel élève : nom, date de naissance, parents, adresse, photo', 'Direction 1, Direction 2'),
    ('Modifier un dossier élève',        'Mettre à jour les informations, changer de classe, modifier les parents', 'Direction 1, Direction 2'),
    ('Supprimer un dossier élève',       'Archiver ou supprimer définitivement un élève du système', 'Direction 1'),
    ('Carte scolaire (QR code)',         'Générer la carte d\'identité scolaire avec photo, nom, classe et QR code unique', 'Direction 1, Direction 2'),
    ('Export toutes les cartes (ZIP)',   'Télécharger l\'ensemble des cartes scolaires en lot pour impression', 'Direction 1, Direction 2'),
    ('Mes enfants (vue parent)',         'Consulter le profil, la classe et le résumé de présence de ses enfants inscrits', 'Parent'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 3. GESTION DES CLASSES
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '3. Gestion des classes', level=1)
add_service_table(doc, [
    ('Liste des classes',           'Voir toutes les classes, les effectifs, les titulaires assignés', 'Direction 1, Direction 2'),
    ('Créer / modifier une classe', 'Définir le nom, le cycle, l\'enseignant titulaire FR et EN', 'Direction 1, Direction 2'),
    ('Supprimer une classe',        'Retirer une classe de l\'année scolaire en cours', 'Direction 1'),
    ('Matières par classe',         'Consulter et gérer les matières enseignées dans chaque classe', 'Direction 1, Direction 2, Enseignant'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 4. GESTION DES COMPTES UTILISATEURS
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '4. Gestion des comptes utilisateurs', level=1)
add_service_table(doc, [
    ('Liste enseignants',              'Voir tous les professeurs, leurs classes assignées et leurs informations', 'Direction 1, Direction 2'),
    ('Liste parents/tuteurs',          'Voir tous les parents inscrits, leurs coordonnées et les enfants associés', 'Direction 1, Direction 2'),
    ('Liste gardiens',                 'Voir le personnel de sécurité et leur activité', 'Direction 1'),
    ('Créer un compte utilisateur',    'Ouvrir un accès pour un enseignant, parent, gardien ou directeur avec photo', 'Direction 1'),
    ('Modifier / supprimer un compte', 'Mettre à jour les informations ou révoquer l\'accès d\'un utilisateur', 'Direction 1'),
    ('Accès parents (contrôle)',       'Définir quels parents peuvent se connecter et ce qu\'ils peuvent consulter', 'Direction 1'),
    ('Photo de profil',                'Chaque utilisateur peut changer sa propre photo de profil', 'Tous sauf élèves'),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PRÉSENCES & ABSENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '5. Présences & absences', level=1)
add_body(doc, 'Suivi en temps réel des présences via le portail QR. Workflow complet de justification.')
add_service_table(doc, [
    ('Tableau des présences',        'Vue journalière par élève et par classe avec statut (présent, retard, excusé, malade, absent)', 'Direction 1, Direction 2, Enseignant, Parent'),
    ('Marquer les présences',        'L\'enseignant valide la liste des scans avant 9h, enregistre les arrivées tardives', 'Enseignant'),
    ('Justifier une absence',        'Le parent soumet une justification écrite (médicale, familiale, autre)', 'Parent'),
    ('Approuver une justification',  'La direction valide ou rejette la justification avec commentaire', 'Direction 1, Direction 2'),
    ('Déclarer sa propre absence',   'Un enseignant déclare son absence (notifie automatiquement la direction)', 'Enseignant'),
    ('Absences de l\'enseignant',    'Historique des absences déclarées par l\'enseignant', 'Enseignant'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 6. PORTAIL DE SÉCURITÉ
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '6. Portail de sécurité (entrées & sorties)', level=1)
add_body(doc, 'Contrôle d\'accès par QR code. Chaque élève possède une carte avec QR unique. '
              'Le gardien scanne à l\'entrée et à la sortie.')
add_service_table(doc, [
    ('Scanner entrée (QR code)',         'Lecture du QR élève à l\'entrée, enregistrement automatique de l\'heure, blocage si frais non payés', 'Gardien'),
    ('Scanner sortie',                   'Identification de l\'élève à la sortie + vérification de la personne venue le récupérer', 'Gardien'),
    ('Personnes autorisées',             'Gérer la liste des personnes autorisées à récupérer chaque élève (photo + relation)', 'Direction 1, Parent'),
    ('Journal des scans',                'Historique complet de toutes les entrées/sorties avec horodatage', 'Direction 1, Direction 2, Gardien'),
    ('Alerte QR inconnu',                'Signal d\'alerte si un QR non reconnu est scanné 3 fois ou plus (tentative d\'intrusion)', 'Direction 1, Gardien'),
    ('Alerte personne non autorisée',    'Notification si quelqu\'un tente de récupérer un élève sans autorisation', 'Direction 1, Gardien'),
    ('Mode LOCKDOWN (urgence)',          'Verrouillage immédiat de toutes les sorties + notification urgente à tous les parents', 'Direction 1'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 7. NOTES & RÉSULTATS
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '7. Notes & résultats académiques', level=1)
add_service_table(doc, [
    ('Saisie des notes',          'L\'enseignant entre les notes par matière et par trimestre (T1, T2, T3)', 'Enseignant'),
    ('Consultation des notes',    'Voir les notes et moyennes par élève, par matière et par trimestre', 'Direction 1, Direction 2, Enseignant, Parent'),
    ('Palmarès / Top 10',         'Classement des meilleurs élèves par moyenne — à l\'échelle de l\'école et par classe', 'Direction 1, Direction 2, Enseignant, Parent'),
    ('Calcul des moyennes',       'Calcul automatique des moyennes par matière, par trimestre et générale', 'Automatique'),
    ('Palmares parent',           'Le parent consulte le classement et la position de son enfant', 'Parent'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 8. RATTRAPAGES
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '8. Cours de rattrapage', level=1)
add_service_table(doc, [
    ('Signaler un élève pour rattrapage', 'L\'enseignant identifie un élève en difficulté et le signale pour une séance de remédiation', 'Enseignant'),
    ('Approuver & fixer le montant',      'La direction valide la demande et fixe les frais de rattrapage à payer par le parent', 'Direction 1, Direction 2'),
    ('Confirmer le paiement',             'Enregistrement du paiement des frais de rattrapage par le parent', 'Direction 1'),
    ('Planifier la séance',               'Fixer la date et l\'heure de la séance de rattrapage dans le calendrier', 'Direction 1'),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. DEVOIRS
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '9. Devoirs & travaux à domicile', level=1)
add_service_table(doc, [
    ('Publier un devoir',             'L\'enseignant crée un devoir avec titre, catégorie, date limite et classe cible', 'Enseignant'),
    ('Modifier / supprimer un devoir', 'Modifier les détails ou retirer un devoir avant l\'échéance', 'Enseignant'),
    ('Liste des devoirs (enseignant)', 'Tableau récapitulatif des devoirs publiés pour sa classe', 'Enseignant'),
    ('Devoirs à venir (parent)',       'Le parent consulte les devoirs des 7 prochains jours pour son enfant', 'Parent'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 10. CAHIER DE TEXTE
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '10. Cahier de texte numérique', level=1)
add_service_table(doc, [
    ('Rédiger une entrée de cours', 'L\'enseignant consigne le contenu du cours, les notions abordées et le programme prévu', 'Enseignant'),
    ('Modifier une entrée',         'Corriger ou compléter une entrée existante du cahier de texte', 'Enseignant'),
    ('Approbation Direction 2',     'La direction pédagogique valide ou rejette les entrées du cahier de texte', 'Direction 2'),
    ('Vue parent',                  'Le parent consulte ce qui a été enseigné dans la classe de son enfant', 'Parent'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 11. APPROBATIONS & WORKFLOW
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '11. Approbations & workflow de validation', level=1)
add_body(doc, 'File d\'attente unifiée pour toutes les demandes nécessitant une validation de la direction.')
add_service_table(doc, [
    ('File d\'approbations unifiée',      'Tableau centralisant toutes les demandes en attente : absences, retards, matières, convocations, rattrapages', 'Direction 1, Direction 2'),
    ('Approuver une absence',             'Valider ou rejeter une justification d\'absence soumise par un parent', 'Direction 1, Direction 2'),
    ('Approuver un retard grave',         'Autoriser exceptionnellement l\'entrée malgré un retard important', 'Direction 1'),
    ('Approuver une nouvelle matière',    'Valider la demande d\'un enseignant d\'ajouter une matière à sa classe', 'Direction 2'),
    ('Valider les rapports caisse',       'Vérifier et approuver les rapports journaliers soumis par le caissier', 'Direction 1'),
    ('Approuver messages enseignants',    'Valider les messages d\'un enseignant avant leur diffusion aux parents', 'Direction 1'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 12. FINANCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '12. Gestion financière', level=1)
add_service_table(doc, [
    ('Tableau des paiements',          'Suivi des frais scolaires par élève, par trimestre (T1/T2/T3) avec statut payé/impayé', 'Direction 1, Direction 3, Parent'),
    ('Enregistrer un paiement',        'Saisir le paiement des frais d\'un élève avec montant, date et mode de règlement', 'Direction 3'),
    ('Enregistrer une dépense',        'Consigner une dépense scolaire avec catégorie, montant et justificatif', 'Direction 3'),
    ('Rapport journalier caisse',      'Résumé des entrées et sorties financières de la journée', 'Direction 3'),
    ('Solde / balance élève',          'Calculer le montant restant à payer pour un élève sur le trimestre en cours', 'Direction 1, Parent'),
    ('Configuration des frais',        'Définir les montants de scolarité pour chaque trimestre', 'Direction 1'),
    ('Blocage automatique au portail', 'Bloquer l\'entrée d\'un élève dont les frais sont en retard (paramétrable)', 'Automatique'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 13. CANTINE
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '13. Service de cantine', level=1)
add_service_table(doc, [
    ('Gestion des inscrits cantine', 'Liste des élèves abonnés au service de restauration et leurs frais', 'Direction 1, Direction 3'),
    ('Menus journaliers',            'Publication du menu du jour avec les plats proposés', 'Direction 3'),
    ('Présences cantine',            'Enregistrer quels élèves ont déjeuné chaque jour', 'Direction 3'),
    ('Vue parent (cantine)',         'Consulter les menus à venir et la participation de son enfant', 'Parent'),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. CONDUITE & DISCIPLINE
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '14. Conduite & discipline', level=1)
add_service_table(doc, [
    ('Enregistrer la conduite',    'Attribuer une note de comportement (Excellent, Bien, Moyen, Mauvais) avec commentaire', 'Direction 1'),
    ('Sanctions disciplinaires',   'Documenter les mesures disciplinaires (avertissement, exclusion temporaire, etc.)', 'Direction 2'),
    ('Historique de conduite',     'Consulter l\'historique complet du comportement d\'un élève', 'Direction 1, Direction 2'),
    ('Vue honneur (parent)',       'Le parent consulte la note de conduite de son enfant depuis son profil', 'Parent'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 15. CONVOCATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '15. Convocations & réunions parents', level=1)
add_service_table(doc, [
    ('Demander une convocation',      'Initier une demande de rencontre avec le parent (automatique si notes insuffisantes)', 'Direction 2'),
    ('Planifier la réunion',          'Fixer la date et l\'heure du rendez-vous avec le parent', 'Direction 1'),
    ('Marquer la présence du parent', 'Enregistrer si le parent s\'est présenté au rendez-vous', 'Direction 1'),
    ('Vue parent (convocations)',     'Le parent consulte ses convocations en attente et ses rendez-vous planifiés', 'Parent'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 16. COMMUNICATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '16. Communication interne', level=1)
add_service_table(doc, [
    ('Messagerie interne',           'Échanges bidirectionnels entre enseignants, parents et direction', 'Tous'),
    ('Boîte réception / envoyés',    'Organiser les messages reçus et envoyés avec filtres', 'Tous'),
    ('Annonce générale',             'Diffuser un message à l\'ensemble des parents de l\'école', 'Direction 1'),
    ('Notifications système',        'Alertes en temps réel avec badges dans la navigation (absences, paiements, scans, etc.)', 'Tous'),
    ('Chat co-titulaire',            'Discussion directe entre les deux enseignants d\'une même classe', 'Enseignant'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 17. EMPLOI DU TEMPS
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '17. Emploi du temps', level=1)
add_service_table(doc, [
    ('Consulter l\'emploi du temps', 'Affichage hebdomadaire (lun–ven) avec matières, heures et enseignants', 'Direction 1, Direction 2, Enseignant, Parent'),
    ('Créer / modifier un créneau', 'Planifier ou ajuster un cours dans l\'emploi du temps d\'une classe', 'Direction 2'),
    ('Supprimer un créneau',        'Retirer un cours de l\'emploi du temps', 'Direction 2'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 18. DOSSIER MÉDICAL
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '18. Dossier médical des élèves', level=1)
add_service_table(doc, [
    ('Créer un dossier médical',        'Enregistrer les informations de santé d\'un élève (allergies, médicaments, pathologies)', 'Direction 1'),
    ('Enregistrer une visite infirmerie', 'Consigner un incident de santé survenu à l\'école avec le traitement effectué', 'Direction 1'),
    ('Historique médical',              'Consulter toutes les interventions médicales passées', 'Direction 1, Parent'),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 19. RESSOURCES HUMAINES
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '19. Ressources humaines (enseignants)', level=1)
add_service_table(doc, [
    ('Salaires mensuels',          'Enregistrer et consulter les salaires versés aux enseignants chaque mois', 'Direction 1'),
    ('Avances sur salaire',        'Gérer les avances demandées par les enseignants', 'Direction 1'),
    ('Primes & bonus',             'Attribuer et consulter les primes d\'incentive pour les enseignants', 'Direction 1, Enseignant'),
    ('Évaluation des enseignants', 'Conduire et archiver les évaluations annuelles de performance', 'Direction 1, Direction 2'),
    ('Activité & connexions',      'Suivre les heures de connexion et les actions effectuées par enseignant', 'Direction 1'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 20. ÉVALUATIONS OFFICIELLES
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '20. Évaluations officielles & conformité', level=1)
add_service_table(doc, [
    ('TENAFEP (CM2)',          'Suivi et dossiers des élèves de 5ème année pour l\'examen national', 'Direction 1, Direction 2'),
    ('Rapport SECOPE/MEPSP',  'Génération du rapport de conformité pour le Ministère de l\'EPST', 'Direction 1, Direction 2'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 21. CALENDRIER SCOLAIRE
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '21. Calendrier scolaire & événements', level=1)
add_service_table(doc, [
    ('Consulter le calendrier', 'Visualiser les événements scolaires à venir (examens, fêtes, journées sportives)', 'Direction 1, Parent'),
    ('Créer / modifier un événement', 'Ajouter ou modifier un événement avec date, titre et description', 'Direction 1'),
    ('Supprimer un événement',       'Retirer un événement du calendrier', 'Direction 1'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 22. ACTIVITÉS PARASCOLAIRES
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '22. Activités parascolaires', level=1)
add_service_table(doc, [
    ('Liste des activités',        'Voir les clubs et activités disponibles (musique, sport, arts, etc.)', 'Direction 1, Parent'),
    ('Inscrire un enfant',         'Le parent inscrit son enfant à une activité parascolaire', 'Parent'),
    ('Créer / modifier une activité', 'Définir une nouvelle activité avec nom, description et capacité', 'Direction 1'),
    ('Gérer les inscriptions',     'Voir la liste des élèves inscrits par activité', 'Direction 1'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 23. EXPORTATION & ARCHIVAGE
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '23. Exportation & archivage', level=1)
add_service_table(doc, [
    ('Export de données',           'Exporter élèves, notes, présences et paiements en CSV ou PDF', 'Direction 1'),
    ('Archiver l\'année scolaire',  'Générer un PDF complet de l\'année puis nettoyer la base pour la rentrée suivante', 'Direction 1'),
    ('Import CSV',                  'Importer en masse des données élèves ou utilisateurs depuis un tableur', 'Direction 1'),
    ('Nettoyage automatique',       'Suppression programmée des logs et notifications de plus de 60 jours', 'Automatique'),
    ('Recherche globale',           'Trouver instantanément n\'importe quel élève, utilisateur ou enregistrement', 'Direction 1'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 24. PROFILS PERSONNELS
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '24. Profils & paramètres personnels', level=1)
add_service_table(doc, [
    ('Modifier ses informations',      'Mettre à jour son numéro de téléphone, son nom et ses préférences', 'Tous'),
    ('Changer sa photo de profil',     'Télécharger ou remplacer sa photo de profil (avec compression automatique)', 'Tous sauf élèves'),
    ('Changer son mot de passe',       'Modifier son mot de passe de connexion de manière sécurisée', 'Tous'),
    ('Statistiques personnelles',      'Tableau de bord de ses propres indicateurs (connexions, actions, résumés)', 'Direction 1, Direction 2'),
    ('Langue préférée (enseignant)',   'Basculer entre français et anglais pour l\'interface', 'Enseignant'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 25. PARAMÈTRES SYSTÈME
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '25. Paramètres système & configuration', level=1)
add_service_table(doc, [
    ('Informations école',         'Configurer le nom, l\'adresse, les contacts et le logo de l\'établissement', 'Direction 1'),
    ('Année scolaire & trimestre', 'Définir l\'année, le trimestre courant et les cycles scolaires', 'Direction 1'),
    ('Horaires scolaires',         'Fixer les heures d\'ouverture, de fermeture et de pause déjeuner', 'Direction 1'),
    ('Toggles système',            'Activer/désactiver les fonctionnalités (alertes, exceptions week-end, lockdown)', 'Direction 1'),
    ('Timeout de session',         'Configurer la durée d\'inactivité avant déconnexion automatique', 'Direction 1'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 26. AUDIT & TRAÇABILITÉ
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '26. Audit & traçabilité', level=1)
add_service_table(doc, [
    ('Journal d\'audit',            'Historique immuable de toutes les actions effectuées dans le système (qui, quoi, quand)', 'Direction 1'),
    ('Contrôle accès parents',      'Décider quels parents peuvent se connecter et limiter leur périmètre de consultation', 'Direction 1'),
])

# ══════════════════════════════════════════════════════════════════════════════
# 27. INFRASTRUCTURE & OFFLINE
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, '27. Infrastructure & disponibilité', level=1)
add_service_table(doc, [
    ('Mode hors-ligne',         'Utilisation complète de l\'application sans connexion internet (données chiffrées AES-256-GCM)', 'Tous'),
    ('Synchronisation auto',    'Envoi automatique des modifications locales vers Supabase dès le retour en ligne', 'Tous'),
    ('Synchronisation manuelle', 'Forcer l\'envoi immédiat des données en attente', 'Tous'),
    ('Application installable (PWA)', 'Installer SchoolSafe comme une application native sur téléphone ou tablette', 'Tous'),
    ('Bilingue FR / EN',        'Interface en français et en anglais selon la préférence de l\'enseignant', 'Enseignant'),
    ('Sécurité des données',    'Données chiffrées en local, accès protégé par RLS Supabase, mots de passe hachés', 'Infrastructure'),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# RÉCAPITULATIF PAR RÔLE
# ══════════════════════════════════════════════════════════════════════════════
add_heading_block(doc, 'Récapitulatif par rôle', level=1)
add_body(doc, 'Résumé des services accessibles selon le profil de connexion.')
doc.add_paragraph()

summary_data = [
    ('Direction 1 (Admin général)',  '27 catégories de services — accès complet à l\'ensemble du système'),
    ('Direction 2 (Pédagogie)',      '~18 catégories — académique, approbations, ressources humaines, sans accès finance'),
    ('Direction 3 (Caisse)',         '~5 catégories — finances, cantine, rapports journaliers, paiements uniquement'),
    ('Enseignant',                   '~9 catégories — sa classe, notes, devoirs, présences, cahier de texte, messagerie, profil'),
    ('Parent',                       '~12 catégories — suivi enfant, notes, présences, devoirs, frais, convocations, messagerie'),
    ('Gardien',                      '~4 catégories — portail entrée/sortie, journal scans, personnes autorisées, profil'),
]
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
hdr = table.rows[0]
for cell, txt in zip(hdr.cells, ['Rôle', 'Périmètre de services']):
    set_cell_bg(cell, '1E3A5F')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
for idx, (role, desc) in enumerate(summary_data):
    row = table.add_row()
    bg  = 'E8F0FB' if idx % 2 == 0 else 'FFFFFF'
    for i, (cell, val) in enumerate(zip(row.cells, [role, desc])):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9.5); r.font.color.rgb = GRAY_TEXT
        if i == 0: r.bold = True
set_table_border(table)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Document généré automatiquement par SchoolSafe v3.0 —')
r.font.size = Pt(8.5); r.font.color.rgb = GRAY_TEXT; r.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
path = '/home/user/zalavrai/SchoolSafe_Services.docx'
doc.save(path)
print(f'Saved: {path}')
