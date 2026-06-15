from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21); s.page_height = Cm(29.7)
s.left_margin = Cm(2.8); s.right_margin = Cm(2.5)
s.top_margin  = Cm(2.5); s.bottom_margin = Cm(2.5)

BLEU   = RGBColor(0x0A, 0x2A, 0x5C)
GOLD   = RGBColor(0xC0, 0x8C, 0x20)
BLANC  = RGBColor(0xFF, 0xFF, 0xFF)
GRIS   = RGBColor(0x44, 0x44, 0x44)
VERT   = RGBColor(0x14, 0x6B, 0x3A)
ROUGE  = RGBColor(0xB0, 0x1C, 0x1C)

ROLE_COLORS = {
    'enseignant': ('1A5276', 'D6EAF8'),
    'direction1': ('145A32', 'D5F5E3'),
    'direction2': ('6E2F89', 'F5E6FF'),
    'caisse':     ('7D3C06', 'FDEBD0'),
    'parent':     ('B7950B', 'FEF9E7'),
    'gardien':    ('922B21', 'FDEDEC'),
}

def set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def spacing(p, line='276'):
    pPr = p._p.get_or_add_pPr()
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:line'), line)
    spc.set(qn('w:lineRule'), 'auto')
    pPr.append(spc)

def h_role(role_name, emoji, subtitle, dark, light):
    """Grand titre de section rôle."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_bg(cell, dark)
    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(emoji + '  ' + role_name.upper())
    r1.bold = True; r1.font.size = Pt(20); r1.font.color.rgb = BLANC
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0xDD,0xDD,0xDD)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def h1(text, sb=18, sa=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(16); run.font.color.rgb = BLEU
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'3');    b.set(qn('w:color'),'0A2A5C')
    pBdr.append(b); pPr.append(pBdr)

def func_block(num, titre, texte, dark='1A5276', light='D6EAF8'):
    """Bloc numéroté pour une fonctionnalité."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Badge numéro
    c0 = tbl.rows[0].cells[0]
    c0.width = Cm(1.4)
    set_bg(c0, dark)
    p0 = c0.add_paragraph(str(num))
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.runs[0].bold = True; p0.runs[0].font.size = Pt(15)
    p0.runs[0].font.color.rgb = BLANC
    # Contenu
    c1 = tbl.rows[0].cells[1]
    set_bg(c1, light)
    pt = c1.add_paragraph()
    rt = pt.add_run('▶  ' + titre)
    rt.bold = True; rt.font.size = Pt(11); rt.font.color.rgb = RGBColor(int(dark[:2],16), int(dark[2:4],16), int(dark[4:],16))
    pb = c1.add_paragraph(texte)
    pb.runs[0].font.size = Pt(10)
    pb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    spacing(pb, '268')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def body(text, sb=3, sa=6, indent=0, size=10.5, color=None, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    spacing(p)
    run = p.add_run(text)
    run.font.size = Pt(size); run.italic = italic; run.bold = bold
    if color: run.font.color.rgb = color

def recap_box(items, dark, light):
    """Encadré récapitulatif des fonctionnalités."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_bg(cell, light)
    pt = cell.add_paragraph()
    rt = pt.add_run('📋  RÉCAPITULATIF DES FONCTIONNALITÉS')
    rt.bold = True; rt.font.size = Pt(10.5)
    rt.font.color.rgb = RGBColor(int(dark[:2],16), int(dark[2:4],16), int(dark[4:],16))
    for item in items:
        pb = cell.add_paragraph('✔  ' + item)
        pb.runs[0].font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def set_bg_simple(cell, hex_color):
    set_bg(cell, hex_color)

# ══════════════════════════════════════════════════════
# PAGE DE GARDE
# ══════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(50)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SCHOOLSAFE v3.0')
r.bold = True; r.font.size = Pt(34); r.font.color.rgb = BLEU

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Guide Complet des Profils Utilisateurs')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = GOLD

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Toutes les fonctionnalités, profil par profil, expliquées en détail')
r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GRIS

doc.add_paragraph()
doc.add_paragraph()

tbl = doc.add_table(rows=5, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
for i,(k,v) in enumerate([
    ('École',    'Complexe Scolaire Le Sage — The Wise School International'),
    ('Ville',    'Kinshasa, République Démocratique du Congo'),
    ('Objet',    'Description détaillée de chaque profil utilisateur'),
    ('Version',  'SchoolSafe v3.0'),
    ('Date',     datetime.date.today().strftime('%d %B %Y')),
]):
    row = tbl.rows[i]
    row.cells[0].text = k; row.cells[1].text = v
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    set_bg(row.cells[0], '0A2A5C')
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = BLANC

doc.add_page_break()

# ══════════════════════════════════════════════════════
# SOMMAIRE
# ══════════════════════════════════════════════════════
h1('SOMMAIRE', sb=0)
for emoji, role, nb in [
    ('👨‍🏫', 'Profil Enseignant',                  '11 fonctionnalités'),
    ('🏫',  'Profil Direction Générale (Dir. 1)',   '14 fonctionnalités'),
    ('📚',  'Profil Direction Pédagogique (Dir. 2)','9 fonctionnalités'),
    ('💰',  'Profil Caisse (Direction 3)',           '7 fonctionnalités'),
    ('👨‍👩‍👧', 'Profil Parent / Tuteur',               '10 fonctionnalités'),
    ('🛡️',  'Profil Gardien de Sécurité',           '7 fonctionnalités'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(emoji + '  '); r1.font.size = Pt(12)
    r2 = p.add_run(role); r2.font.size = Pt(11); r2.bold = True
    r3 = p.add_run('  —  ' + nb); r3.font.size = Pt(10); r3.font.color.rgb = GOLD

doc.add_page_break()

# ══════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════
h1('Introduction')
body('''SchoolSafe v3.0 est organisé autour de six profils distincts, chacun conçu pour correspondre exactement au rôle d\'un acteur spécifique au sein du Complexe Scolaire Le Sage. Chaque profil possède son propre tableau de bord, ses propres fonctionnalités et ses propres droits d\'accès. Ce document décrit, profil par profil et fonctionnalité par fonctionnalité, ce que chaque utilisateur peut faire dans l\'application — avec des explications claires et détaillées pour chacune de ces fonctionnalités.''')
body('''Les six profils sont : l\'Enseignant, la Direction Générale (Direction 1), la Direction Pédagogique (Direction 2), la Caisse (Direction 3), le Parent ou Tuteur, et le Gardien de Sécurité.''')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# PROFIL 1 : ENSEIGNANT
# ══════════════════════════════════════════════════════
D, L = '1A5276', 'D6EAF8'
h_role('PROFIL ENSEIGNANT', '👨‍🏫', 'Le professeur — au cœur de la pédagogie', D, L)

body('''L\'enseignant est le principal producteur d\'informations pédagogiques dans SchoolSafe. Chaque enseignant est assigné à une ou plusieurs classes spécifiques. Il n\'a accès qu\'aux données de ses propres classes. Son interface est disponible en français ou en anglais selon la section à laquelle il est rattaché. Voici les 11 fonctionnalités dont il dispose.''')

func_block(1, 'TABLEAU DE BORD DE SA CLASSE',
'''Dès la connexion, l\'enseignant voit son tableau de bord personnalisé. Ce tableau affiche un résumé visuel de sa classe : le nombre total d\'élèves, le nombre de présents du jour, la dernière note saisie, les devoirs récemment publiés et les éventuelles notifications non lues. C\'est le point de départ de toutes ses actions quotidiennes dans l\'application.''', D, L)

func_block(2, 'GESTION DES PRÉSENCES',
'''C\'est la tâche quotidienne numéro un de l\'enseignant. Chaque jour, il ouvre la section "Présences", et pour chaque élève de sa classe, il sélectionne l\'un des trois statuts : Présent, Absent ou En retard. Pour les élèves en retard, il peut saisir l\'heure exacte d\'arrivée. Une fois enregistrées, les présences sont immédiatement visibles par la Direction 2 et par les parents des élèves concernés. Si un élève est absent, le parent reçoit automatiquement une notification sur son téléphone. L\'enseignant ne peut pas modifier une présence après sa validation par la direction.''', D, L)

func_block(3, 'SAISIE DES NOTES PAR MATIÈRE',
'''L\'enseignant entre les notes de ses élèves dans SchoolSafe par matière, par type d\'évaluation (devoir sur table, interrogation, examen trimestriel) et par trimestre (T1, T2 ou T3). Il saisit une note pour chaque élève de sa classe. L\'application calcule automatiquement la moyenne de la classe et met à jour les statistiques générales. Les notes restent en état "brouillon" tant qu\'elles ne sont pas soumises à la Direction 2 : ni la direction ni les parents ne les voient dans cet état. L\'enseignant peut modifier librement les notes en brouillon avant de les soumettre.''', D, L)

func_block(4, 'SOUMISSION DES GRILLES DE NOTES À LA DIRECTION',
'''Lorsque toutes les notes d\'une évaluation sont saisies et vérifiées, l\'enseignant les soumet officiellement à la Direction Pédagogique (Direction 2) pour validation. Cette action déclenche une notification à la direction et place la grille en état "en attente de validation". L\'enseignant ne peut plus modifier les notes après soumission, sauf si la direction les lui renvoie avec des corrections à apporter. Une fois la grille validée par la Direction 2, les notes deviennent officielles et apparaissent automatiquement chez les parents.''', D, L)

func_block(5, 'PUBLICATION DES DEVOIRS',
'''L\'enseignant publie les devoirs et exercices donnés aux élèves directement dans l\'application. Il renseigne le titre du devoir, la matière concernée, la catégorie (devoir à la maison, exposé, révision, travail de groupe) et la date limite de remise. Une fois publié, le devoir est immédiatement visible par tous les parents des élèves de la classe. L\'enseignant peut modifier un devoir publié si une erreur s\'est glissée, ou le supprimer s\'il est annulé. L\'historique des devoirs publiés reste accessible à tout moment.''', D, L)

func_block(6, 'CAHIER DE TEXTE NUMÉRIQUE',
'''Le cahier de texte est le journal pédagogique de l\'enseignant. Après chaque cours, il renseigne dans SchoolSafe ce qui a été fait : les notions enseignées, les exercices réalisés, les thèmes abordés, les documents utilisés et les travaux prévus pour la séance suivante. Ce document est consultable à tout moment par la Direction Pédagogique, qui l\'utilise pour vérifier que le programme officiel est bien respecté et que l\'enseignant avance au rythme attendu. Le cahier de texte numérique remplace définitivement les cahiers papier souvent perdus ou mal tenus.''', D, L)

func_block(7, 'SIGNALEMENT DES RATTRAPAGES',
'''Lorsqu\'un élève n\'a pas pu participer à une évaluation pour une raison valable (maladie, absence justifiée, force majeure), l\'enseignant peut signaler dans SchoolSafe qu\'un rattrapage est nécessaire. Il sélectionne l\'élève concerné, précise la matière et l\'évaluation manquée, et indique le motif de la demande. Cette demande est transmise automatiquement à la Direction 2, qui l\'examine et décide de l\'approuver ou de la refuser. L\'enseignant est notifié de la décision. Si le rattrapage est approuvé, il organise l\'évaluation de rattrapage et saisit la note obtenue dans le système, en suivant le même circuit de validation que les notes ordinaires.''', D, L)

func_block(8, 'CONSULTATION DE L\'EMPLOI DU TEMPS',
'''L\'enseignant peut consulter son emploi du temps hebdomadaire directement dans SchoolSafe. Il voit les cours qu\'il doit dispenser, dans quelle classe, à quelle heure et quel jour. Cette fonctionnalité lui permet d\'organiser sa semaine et de planifier ses évaluations en conséquence.''', D, L)

func_block(9, 'CONSULTATION ET SUIVI DES APPROBATIONS',
'''L\'enseignant peut soumettre des demandes d\'approbation à la direction pour des décisions qui nécessitent une validation hiérarchique. Il consulte l\'état de ses demandes en cours (en attente, approuvée, refusée) et reçoit une notification dès qu\'une décision est prise. Ce système formalise les échanges entre enseignant et direction et crée une trace numérique de toutes les décisions.''', D, L)

func_block(10, 'RÉCEPTION DES NOTIFICATIONS',
'''L\'enseignant reçoit des notifications officielles dans SchoolSafe. Ces notifications peuvent venir de la Direction 1 (convocations, instructions générales), de la Direction 2 (retour sur une grille de notes, demande de correction, rappel de délai) ou être générées automatiquement par le système (grille de notes validée, rattrapage approuvé ou refusé). Un badge de comptage visible dès l\'ouverture de l\'application indique le nombre de notifications non lues.''', D, L)

func_block(11, 'GESTION DU PROFIL ET SÉCURITÉ DU COMPTE',
'''L\'enseignant gère ses informations personnelles depuis la section "Profil". Il peut mettre à jour son numéro de téléphone, qui sert de contact officiel dans l\'école, et modifier son code PIN de connexion à tout moment. Il est responsable de la confidentialité de son code PIN. Toutes les actions effectuées depuis son compte sont enregistrées dans le journal d\'audit et lui sont directement attribuées. En cas d\'anomalie ou de suspicion de compromission de son compte, il doit contacter la Direction 1 immédiatement.''', D, L)

recap_box([
    'Tableau de bord de sa classe',
    'Gestion des présences (présent / absent / en retard + heure)',
    'Saisie des notes par matière, type et trimestre',
    'Soumission des grilles de notes à la Direction 2',
    'Publication des devoirs',
    'Cahier de texte numérique',
    'Signalement des rattrapages',
    'Consultation de l\'emploi du temps',
    'Suivi des approbations',
    'Réception des notifications',
    'Gestion du profil et sécurité du compte',
], D, L)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# PROFIL 2 : DIRECTION 1
# ══════════════════════════════════════════════════════
D, L = '145A32', 'D5F5E3'
h_role('DIRECTION GÉNÉRALE — DIRECTION 1', '🏫', 'Administrateur général — accès complet à toute l\'école', D, L)

body('''La Direction 1 est l\'administrateur général de SchoolSafe. Elle dispose du niveau d\'accès le plus élevé dans le système et peut consulter, modifier et superviser l\'intégralité des données de l\'école. Elle est responsable de la configuration du système, de la gestion des utilisateurs et de la gouvernance globale de l\'établissement. Voici les 14 fonctionnalités dont elle dispose.''')

func_block(1, 'TABLEAU DE BORD GÉNÉRAL DE L\'ÉCOLE',
'''Le tableau de bord de la Direction 1 offre une vue panoramique de toute l\'école en temps réel. Il affiche le nombre total d\'élèves inscrits, le taux de présence du jour, le nombre de paiements enregistrés ce mois-ci, le montant total des frais collectés, les alertes urgentes en cours (lockdown, notifications non traitées) et le top 10 des meilleurs élèves de l\'école. Ce tableau de bord est le centre de commandement de la direction.''', D, L)

func_block(2, 'GESTION COMPLÈTE DES ÉLÈVES',
'''La Direction 1 gère l\'intégralité des dossiers élèves. Elle peut créer un nouveau dossier élève en saisissant son nom, prénom, numéro de matricule, date de naissance, classe d\'affectation, identifiant du parent et photo. Elle peut modifier un dossier existant, corriger des erreurs ou mettre à jour les informations. Elle peut bloquer un élève — empêchant son entrée via le portail QR Code — pour des raisons disciplinaires ou administratives (frais impayés). Elle peut le débloquer à tout moment. Elle peut également supprimer définitivement un dossier élève si nécessaire.''', D, L)

func_block(3, 'GÉNÉRATION DES CARTES QR CODE',
'''La Direction 1 génère les cartes d\'identité numérique de chaque élève. Ces cartes contiennent la photo de l\'élève, son nom complet, sa classe, son matricule, le logo de l\'école et son QR Code unique. Elle peut générer une carte individuelle pour un élève spécifique ou exporter en une seule opération l\'ensemble des cartes de tous les élèves de l\'école dans un fichier ZIP, prêt à être envoyé à un imprimeur pour plastification.''', D, L)

func_block(4, 'GESTION DES CLASSES',
'''La Direction 1 crée et configure les classes de l\'école. Pour chaque classe, elle définit son nom, son cycle (Maternelle, Primaire ou Secondaire) et assigne les enseignants titulaires — un pour la section francophone et un pour la section anglophone si nécessaire. Elle peut consulter la liste complète des élèves de chaque classe et modifier la configuration à tout moment.''', D, L)

func_block(5, 'GESTION DE TOUS LES COMPTES UTILISATEURS',
'''La Direction 1 est la seule à pouvoir créer des comptes pour tous les autres rôles du système : enseignants, parents, gardiens et agent de caisse. Pour chaque compte, elle définit le nom, le rôle, le numéro de téléphone et le code d\'accès initial. Elle peut modifier les informations d\'un compte existant, réinitialiser le code PIN d\'un utilisateur qui l\'a oublié, et supprimer un compte lorsqu\'un membre du personnel quitte l\'école ou qu\'un parent n\'a plus d\'enfant dans l\'établissement.''', D, L)

func_block(6, 'GESTION DES ACCÈS PARENTS AUX ÉLÈVES',
'''La Direction 1 gère les liens entre les parents et leurs enfants dans le système. Elle associe chaque parent aux élèves dont il a la charge. Si un parent a plusieurs enfants dans l\'école, elle les relie tous à son compte. Elle peut également retirer l\'accès d\'un parent à un élève si la situation familiale change.''', D, L)

func_block(7, 'PARAMÈTRES DE L\'ÉCOLE',
'''La Direction 1 configure les paramètres généraux de l\'école. Elle définit les frais scolaires par trimestre (T1, T2, T3) et peut les modifier en début d\'année ou en cas de changement. Elle choisit le trimestre en cours, ce qui oriente les saisies de notes et de paiements vers la bonne période. Elle active ou désactive des fonctionnalités spécifiques de l\'application selon les besoins de l\'école.''', D, L)

func_block(8, 'MODE LOCKDOWN — CONFINEMENT D\'URGENCE',
'''La Direction 1 peut activer en un seul clic le mode Lockdown, un système de confinement d\'urgence. Dès son activation, toutes les sorties d\'élèves sont bloquées. Le gardien voit immédiatement une alerte rouge clignotante sur son écran lui indiquant qu\'aucun élève ne peut quitter l\'école. Ce mode peut être déclenché en cas d\'alerte sécuritaire, d\'incident grave ou de toute situation nécessitant que tous les élèves restent à l\'intérieur. Seule la Direction 1 peut désactiver le lockdown.''', D, L)

func_block(9, 'CONVOCATIONS',
'''La Direction 1 peut envoyer des convocations officielles à n\'importe quel utilisateur du système : parents, enseignants, gardien. La convocation est envoyée sous forme de notification dans l\'application et contient le message officiel de la direction. Toutes les convocations envoyées sont conservées dans l\'historique.''', D, L)

func_block(10, 'RAPPORT D\'ABSENCES ET DE PRÉSENCES',
'''La Direction 1 accède aux rapports détaillés d\'absences et de présences pour toute l\'école. Elle peut filtrer par classe, par date, par élève ou par période. Elle peut identifier les élèves avec un taux d\'absentéisme préoccupant et décider de convoquer leurs parents. Ces rapports peuvent être consultés à l\'écran ou exportés.''', D, L)

func_block(11, 'CALENDRIER SCOLAIRE',
'''La Direction 1 gère le calendrier officiel de l\'école dans SchoolSafe. Elle peut y inscrire les événements importants : jours fériés, examens, réunions de parents, cérémonies, congés. Ce calendrier est visible par les enseignants depuis leur module.''', D, L)

func_block(12, 'ARCHIVAGE ANNUEL DE FIN D\'ANNÉE',
'''En fin d\'année scolaire, la Direction 1 procède à l\'archivage complet. D\'un seul clic, SchoolSafe génère un document PDF officiel contenant l\'intégralité des données de l\'année écoulée : liste complète des élèves avec leurs photos, résultats scolaires par matière et par trimestre, palmares de chaque classe, états de paiement, journal des présences et historique des scans d\'entrée/sortie. Ce PDF est produit avec le logo de l\'école et sert d\'archive officielle définitive. Une fois validé, les données de l\'année sont nettoyées pour préparer la rentrée suivante.''', D, L)

func_block(13, 'JOURNAL D\'AUDIT',
'''Le journal d\'audit est un enregistrement chronologique et infalsifiable de toutes les actions effectuées dans SchoolSafe par n\'importe quel utilisateur. Chaque entrée indique : la date et l\'heure exacte, le nom de l\'utilisateur, la nature de l\'action effectuée et les données concernées. La Direction 1 est la seule à avoir accès à ce journal. Elle peut l\'utiliser pour détecter des anomalies, retrouver l\'origine d\'une erreur, ou identifier une tentative de falsification des données. Les entrées de plus de 60 jours sont automatiquement nettoyées.''', D, L)

func_block(14, 'PROFIL, STATISTIQUES ET SÉCURITÉ',
'''La Direction 1 dispose d\'une section profil complète avec des statistiques de haut niveau sur l\'ensemble de l\'école. Elle peut y modifier ses informations personnelles (téléphone), changer son mot de passe, et consulter son historique d\'activité. Elle voit également des données agrégées : nombre total d\'élèves par cycle, taux global de présence, montant total des frais collectés par rapport aux frais dus.''', D, L)

recap_box([
    'Tableau de bord général de l\'école',
    'Gestion complète des élèves (créer, modifier, bloquer, supprimer)',
    'Génération des cartes QR Code (individuelle ou ZIP complet)',
    'Gestion des classes et affectation des enseignants',
    'Gestion de tous les comptes utilisateurs (tous rôles)',
    'Gestion des accès parents aux élèves',
    'Paramètres de l\'école (frais, trimestre, toggles)',
    'Mode Lockdown — confinement d\'urgence',
    'Envoi de convocations officielles',
    'Rapports d\'absences et de présences',
    'Calendrier scolaire',
    'Archivage annuel + génération du PDF officiel',
    'Journal d\'audit complet',
    'Profil, statistiques globales et sécurité',
], D, L)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# PROFIL 3 : DIRECTION 2
# ══════════════════════════════════════════════════════
D, L = '6E2F89', 'F5E6FF'
h_role('DIRECTION PÉDAGOGIQUE — DIRECTION 2', '📚', 'Responsable pédagogique — supervision de l\'enseignement', D, L)

body('''La Direction 2 est le responsable pédagogique de l\'école. Elle supervise l\'ensemble du travail des enseignants, valide les notes, gère les rattrapages et s\'assure de la qualité de l\'enseignement dans chaque classe. Elle a accès aux données de toutes les classes mais ne peut pas gérer les utilisateurs ni les paramètres de l\'école — ce sont les prérogatives de la Direction 1. Voici les 9 fonctionnalités dont elle dispose.''')

func_block(1, 'TABLEAU DE BORD PÉDAGOGIQUE',
'''Le tableau de bord de la Direction 2 lui donne une vue d\'ensemble pédagogique de l\'école. Elle voit le nombre de grilles de notes en attente de validation, les rattrapages en attente de décision, le taux de remplissage des cahiers de texte par les enseignants et les statistiques de présences globales. Ce tableau de bord est son outil de pilotage quotidien de la pédagogie.''', D, L)

func_block(2, 'VALIDATION DES GRILLES DE NOTES',
'''La Direction 2 reçoit toutes les grilles de notes soumises par les enseignants. Pour chaque grille, elle peut voir la liste complète des notes de chaque élève, la moyenne de la classe, le type d\'évaluation et le trimestre concerné. Elle vérifie la cohérence des notes, s\'assure qu\'aucune valeur n\'est aberrante et qu\'aucun élève n\'a été oublié. Si la grille est correcte, elle la valide : les notes deviennent officielles et apparaissent chez les parents. Si elle détecte un problème, elle la renvoie à l\'enseignant avec une explication, et l\'enseignant reçoit une notification pour corriger et resoumettre.''', D, L)

func_block(3, 'PALMARES ET CLASSEMENTS',
'''La Direction 2 accède au palmares de chaque classe pour chaque trimestre. Le palmares est le classement officiel des élèves par ordre de mérite, basé sur leur moyenne générale. Elle peut consulter le top 10 des élèves de toute l\'école, identifier les élèves en tête de classe et ceux en difficulté, et comparer les performances d\'une classe à l\'autre ou d\'un trimestre à l\'autre.''', D, L)

func_block(4, 'SUPERVISION DES PRÉSENCES ET ABSENCES',
'''La Direction 2 consulte les rapports de présences de toutes les classes. Elle peut filtrer par date, par classe ou par élève. Elle identifie les élèves avec un taux d\'absentéisme préoccupant et peut signaler la situation à la Direction 1 pour qu\'une convocation soit envoyée aux parents. Elle vérifie également que les enseignants marquent bien les présences régulièrement.''', D, L)

func_block(5, 'GESTION ET VALIDATION DES RATTRAPAGES',
'''Toutes les demandes de rattrapage signalées par les enseignants arrivent dans le tableau de bord de la Direction 2. Elle examine chaque demande individuellement : elle vérifie les antécédents d\'absence de l\'élève dans SchoolSafe, évalue si le motif invoqué est recevable et décide d\'approuver ou de refuser le rattrapage. Sa décision est enregistrée avec horodatage. L\'enseignant et la Direction 1 sont informés de la décision par notification.''', D, L)

func_block(6, 'GESTION DES APPROBATIONS',
'''La Direction 2 traite les demandes d\'approbation soumises par les enseignants pour toute décision nécessitant une validation hiérarchique. Elle peut approuver ou refuser chaque demande avec un commentaire explicatif. L\'historique de toutes les approbations traitées est conservé dans le système avec la décision prise et sa date.''', D, L)

func_block(7, 'CONSULTATION DU CAHIER DE TEXTE DE TOUS LES ENSEIGNANTS',
'''La Direction 2 accède au cahier de texte numérique de chaque enseignant de l\'école. Elle peut voir, pour chaque enseignant et chaque classe, le détail de tous les cours dispensés : date, matière, contenu pédagogique, exercices réalisés et travaux prévus. Cette fonctionnalité lui permet de vérifier que le programme officiel est respecté, que les enseignants avancent au bon rythme et qu\'aucune partie du programme n\'est négligée.''', D, L)

func_block(8, 'CONSULTATION DE L\'EMPLOI DU TEMPS',
'''La Direction 2 consulte l\'emploi du temps de toutes les classes. Elle s\'assure que la planification est cohérente avec le calendrier scolaire et que chaque classe dispose du nombre d\'heures de cours requis par matière.''', D, L)

func_block(9, 'PROFIL ET SÉCURITÉ',
'''La Direction 2 gère ses informations personnelles depuis sa section profil : mise à jour de son numéro de téléphone et changement de son mot de passe de connexion. Elle dispose également d\'une vue statistique de son activité dans l\'application : nombre de grilles validées, nombre de rattrapages traités, temps moyen de validation.''', D, L)

recap_box([
    'Tableau de bord pédagogique',
    'Validation (ou renvoi) des grilles de notes soumises par les enseignants',
    'Palmares et classements par classe et par trimestre',
    'Supervision des présences et absences de toutes les classes',
    'Gestion et validation des rattrapages',
    'Gestion des approbations',
    'Consultation du cahier de texte de tous les enseignants',
    'Consultation de l\'emploi du temps',
    'Profil et sécurité',
], D, L)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# PROFIL 4 : CAISSE
# ══════════════════════════════════════════════════════
D, L = '7D3C06', 'FDEBD0'
h_role('CAISSE — DIRECTION 3', '💰', 'Agent de caisse — gestion financière de l\'école', D, L)

body('''La Direction 3, ou agent de caisse, est responsable de toute la gestion financière de l\'école dans SchoolSafe. Elle enregistre les paiements des frais scolaires, suit les recettes et dépenses journalières et produit les rapports financiers. Son module est conçu pour être simple, rapide et infalsifiable. Voici les 7 fonctionnalités dont elle dispose.''')

func_block(1, 'TABLEAU DE BORD FINANCIER',
'''Le tableau de bord de l\'agent de caisse affiche un résumé financier en temps réel : le montant total encaissé dans la journée, le nombre de paiements enregistrés ce mois-ci, le solde global de caisse (recettes moins dépenses) et une alerte sur les élèves ayant un solde dû important. Ce tableau de bord lui permet de démarrer chaque journée avec une vision claire de la situation financière de l\'école.''', D, L)

func_block(2, 'ENREGISTREMENT DES PAIEMENTS DE FRAIS SCOLAIRES',
'''C\'est la fonctionnalité principale de l\'agent de caisse. Lorsqu\'un parent vient payer les frais scolaires, l\'agent recherche l\'élève dans SchoolSafe par son nom ou son matricule. Il sélectionne le trimestre concerné (T1, T2 ou T3) et saisit le montant reçu. L\'application calcule instantanément le solde restant dû, prend en compte les paiements partiels déjà effectués et met à jour le compte de l\'élève. Le paiement est enregistré avec la date, l\'heure et le nom de l\'agent. Le parent reçoit une notification confirmant le paiement.''', D, L)

func_block(3, 'CONSULTATION DES SOLDES PAR ÉLÈVE',
'''L\'agent de caisse peut consulter à tout moment l\'état financier de chaque élève : le montant total des frais dus pour l\'année, le montant déjà payé trimestre par trimestre et le solde restant. Il peut également voir l\'historique complet de tous les versements effectués par un parent donné, avec les dates et montants exacts.''', D, L)

func_block(4, 'CAISSE JOURNALIÈRE — RECETTES ET DÉPENSES',
'''Au-delà des frais scolaires, l\'agent de caisse enregistre dans SchoolSafe toutes les entrées et sorties d\'argent de la caisse de l\'école. Pour chaque recette (vente de fournitures, inscription, autres), il saisit le montant et un libellé descriptif. Pour chaque dépense (achat de matériel, réparation, frais divers), il fait de même. L\'application calcule automatiquement le solde de caisse du jour et le cumul depuis le début du mois. La Direction 1 peut consulter ces données à tout moment.''', D, L)

func_block(5, 'LISTE DES IMPAYÉS',
'''SchoolSafe génère automatiquement la liste de tous les élèves ayant un solde dû, c\'est-à-dire dont les frais scolaires ne sont pas intégralement payés. Cette liste est filtrée par classe et par trimestre. Elle permet à la caisse de cibler les familles à relancer et à la Direction 1 de décider des éventuelles mesures à prendre (blocage de l\'accès, convocation).''', D, L)

func_block(6, 'EXPORTS ET RAPPORTS FINANCIERS',
'''L\'agent de caisse peut exporter les données financières en différents formats. Il peut générer un rapport des paiements par période (jour, semaine, mois, trimestre), un rapport des impayés, et un récapitulatif des recettes et dépenses journalières. Ces exports peuvent être utilisés pour la comptabilité mensuelle de l\'école ou pour les rapports destinés à la direction.''', D, L)

func_block(7, 'PROFIL ET SÉCURITÉ',
'''Comme pour tous les autres utilisateurs, l\'agent de caisse gère ses informations personnelles depuis sa section profil. Il peut modifier son numéro de téléphone et changer son code PIN. Toutes ses opérations financières sont enregistrées dans le journal d\'audit avec son nom et l\'heure exacte, garantissant une traçabilité totale de chaque transaction.''', D, L)

recap_box([
    'Tableau de bord financier (recettes du jour, solde caisse)',
    'Enregistrement des paiements de frais scolaires (T1 / T2 / T3)',
    'Consultation des soldes et historique de paiements par élève',
    'Caisse journalière : enregistrement recettes et dépenses',
    'Liste automatique des impayés par classe',
    'Exports et rapports financiers',
    'Profil et sécurité',
], D, L)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# PROFIL 5 : PARENT
# ══════════════════════════════════════════════════════
D, L = 'B7950B', 'FEF9E7'
h_role('PROFIL PARENT / TUTEUR', '👨‍👩‍👧', 'Le parent — informé en temps réel sur la scolarité de son enfant', D, L)

body('''Le profil Parent est conçu pour donner aux familles une visibilité complète et en temps réel sur tout ce qui concerne la scolarité de leurs enfants. L\'accès est en lecture seule : le parent ne peut pas modifier de données, mais il peut consulter tout ce qui le concerne. Si un parent a plusieurs enfants dans l\'école, ils apparaissent tous dans son profil et il peut passer de l\'un à l\'autre facilement. Voici les 10 fonctionnalités dont il dispose.''')

func_block(1, 'TABLEAU DE BORD PAR ENFANT',
'''Le tableau de bord du parent affiche un résumé par enfant : le statut de présence du jour (présent, absent, en retard), la dernière note enregistrée, le nombre de devoirs publiés cette semaine et les notifications non lues. Si le parent a plusieurs enfants, il voit une fiche par enfant et peut naviguer facilement de l\'un à l\'autre.''', D, L)

func_block(2, 'CONSULTATION DES NOTES ET DU BULLETIN',
'''Le parent peut consulter toutes les notes de son enfant au fur et à mesure qu\'elles sont saisies et validées par les enseignants et la direction. Il voit les notes par matière, par type d\'évaluation (devoir, interrogation, examen) et par trimestre. L\'application calcule et affiche automatiquement la moyenne par matière et la moyenne générale. Le parent voit donc en temps réel le niveau scolaire de son enfant, sans attendre la remise du bulletin trimestriel.''', D, L)

func_block(3, 'PALMARES ET RANG DE L\'ENFANT',
'''Le parent peut accéder au palmares de la classe de son enfant, c\'est-à-dire le classement officiel de tous les élèves par ordre de mérite. Il voit la position de son enfant dans le classement, sa moyenne générale et sa progression par rapport aux trimestres précédents. Cette fonctionnalité stimule la motivation des élèves et permet aux parents de situer leurs enfants dans leur environnement scolaire.''', D, L)

func_block(4, 'SUIVI DES PRÉSENCES ET DES ABSENCES',
'''Le parent consulte l\'historique complet des présences de son enfant : les jours de présence, les absences et les retards, avec la date et l\'heure exacte dans le cas des retards. Si son enfant est absent, il reçoit une notification automatique le matin même. Cette fonctionnalité lui permet de s\'assurer que son enfant arrive bien à l\'école et de réagir immédiatement en cas d\'absence inexpliquée.''', D, L)

func_block(5, 'CONSULTATION DES DEVOIRS',
'''Le parent voit tous les devoirs publiés par les enseignants de la classe de son enfant. Pour chaque devoir, il connaît la matière, le titre, la description et la date limite de remise. Cette fonctionnalité lui permet de s\'assurer que son enfant fait bien ses devoirs et de l\'aider si nécessaire.''', D, L)

func_block(6, 'ÉTAT DES PAIEMENTS DES FRAIS SCOLAIRES',
'''Le parent consulte l\'état financier de son dossier dans SchoolSafe : le montant total des frais scolaires de l\'année, ce qui a déjà été payé trimestre par trimestre et ce qui reste dû. Il voit l\'historique de tous ses versements avec les dates et montants. Cette transparence lui permet de planifier ses paiements et de vérifier que ses versements ont bien été enregistrés par la caisse.''', D, L)

func_block(7, 'JOURNAL DES ENTRÉES ET SORTIES (SCAN LOG)',
'''Le parent peut consulter le journal de toutes les entrées et sorties de son enfant enregistrées par le système de scan QR Code à la grille de l\'école. Il voit chaque passage avec la date, l\'heure et le type de mouvement (entrée ou sortie). Cette fonctionnalité lui permet de vérifier que son enfant est bien arrivé à l\'école le matin et qu\'il est bien rentré à l\'heure prévue le soir.''', D, L)

func_block(8, 'CONSULTATION DES PERSONNES AUTORISÉES',
'''Le parent peut consulter la liste des personnes autorisées à récupérer son enfant à la sortie de l\'école, telle qu\'elle est enregistrée dans SchoolSafe. Il voit le nom, la relation et la photo de chaque personne. Si cette liste doit être modifiée (ajout ou retrait d\'une personne), il doit en faire la demande à la direction, qui effectue les modifications.''', D, L)

func_block(9, 'RÉCEPTION DES NOTIFICATIONS',
'''Le parent reçoit des notifications officielles de l\'école directement dans SchoolSafe : convocations de la direction, rappels de paiement, alertes concernant l\'absence de son enfant, informations générales. Ces notifications sont affichées avec un badge de comptage sur l\'icône de l\'application et peuvent être marquées comme lues ou supprimées. L\'historique des notifications est conservé.''', D, L)

func_block(10, 'TABLEAU D\'HONNEUR ET PROFIL',
'''Si son enfant figure parmi les meilleurs élèves du trimestre, le parent voit une mention spéciale "Tableau d\'honneur" dans son profil. Cette reconnaissance officielle encourage les élèves et valorise les efforts des familles. Dans la section profil, le parent peut mettre à jour son numéro de téléphone et changer son code PIN de connexion.''', D, L)

recap_box([
    'Tableau de bord par enfant (présence, notes, devoirs)',
    'Consultation des notes et moyennes en temps réel',
    'Palmares et rang de l\'enfant dans sa classe',
    'Suivi des présences, absences et retards',
    'Consultation des devoirs publiés par les enseignants',
    'État des paiements et historique des versements',
    'Journal des entrées et sorties (scan QR Code)',
    'Consultation des personnes autorisées à la sortie',
    'Réception des notifications officielles de l\'école',
    'Tableau d\'honneur et gestion du profil',
], D, L)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# PROFIL 6 : GARDIEN
# ══════════════════════════════════════════════════════
D, L = '922B21', 'FDEDEC'
h_role('PROFIL GARDIEN DE SÉCURITÉ', '🛡️', 'Premier rempart de la sécurité physique des élèves', D, L)

body('''Le gardien est responsable du contrôle physique de toutes les entrées et sorties de l\'école. Son module dans SchoolSafe est conçu pour être utilisé rapidement, y compris dans des conditions agitées (heure de pointe), par un agent qui n\'est pas nécessairement un expert en informatique. L\'interface est volontairement simple, visuelle et réactive. Voici les 7 fonctionnalités dont il dispose.''')

func_block(1, 'TABLEAU DE BORD DE SÉCURITÉ',
'''Le tableau de bord du gardien affiche les informations essentielles pour sa mission : le nombre d\'élèves entrés dans l\'école depuis le début de la journée, le nombre de sorties enregistrées, le statut du mode Lockdown (actif ou inactif) et les alertes en cours. Si le mode Lockdown est activé par la direction, une alerte rouge clignotante occupe tout l\'écran, avertissant immédiatement le gardien qu\'aucune sortie ne doit être autorisée.''', D, L)

func_block(2, 'SCANNER D\'ENTRÉE (QR CODE)',
'''C\'est la fonctionnalité centrale du gardien le matin. Il pointe la caméra de son téléphone sur la carte QR Code de l\'élève. En moins d\'une seconde, l\'application affiche la photo de l\'élève en grand format, son nom et sa classe, et indique clairement l\'une des trois réponses possibles : fond vert (AUTORISÉ — l\'élève peut entrer normalement), fond orange (ALERTE — situation à vérifier avec la direction), fond rouge (REFUSÉ — l\'élève est bloqué ou l\'école est en lockdown). L\'heure d\'entrée est automatiquement enregistrée dans le journal de l\'école.''', D, L)

func_block(3, 'SCANNER DE SORTIE AVEC VÉRIFICATION DES PERSONNES AUTORISÉES',
'''À la sortie, lorsqu\'une personne se présente pour récupérer un élève, le gardien scanne la carte QR Code de l\'enfant. SchoolSafe affiche immédiatement la liste de toutes les personnes autorisées à récupérer cet élève, avec la photo de chaque personne en format suffisamment grand pour être reconnue visuellement. Le gardien compare la personne devant lui avec les photos affichées. Si le visage correspond à une personne autorisée, il valide la sortie. Si la personne n\'est pas dans la liste, il refuse la sortie et contacte la direction. L\'heure de sortie est enregistrée automatiquement.''', D, L)

func_block(4, 'ZOOM SUR LA PHOTO DE LA PERSONNE AUTORISÉE',
'''Pour faciliter l\'identification visuelle, le gardien peut appuyer sur la photo d\'une personne autorisée pour l\'afficher en plein écran sur son téléphone. Cette fonctionnalité est particulièrement utile en plein soleil, à distance, ou pour des personnes dont les traits peuvent se ressembler. Le gardien dispose ainsi d\'une référence visuelle claire et sans ambiguïté.''', D, L)

func_block(5, 'JOURNAL DES SCANS DE LA JOURNÉE',
'''Le gardien peut consulter à tout moment le journal chronologique de tous les scans de la journée : chaque entrée et chaque sortie enregistrée, avec le nom de l\'élève, l\'heure exacte et le résultat du scan (autorisé, alerte, refusé). Ce journal lui permet de retrouver rapidement un passage spécifique en cas de question ou de litige, et de vérifier qu\'un élève signalé absent est bien absent du journal d\'entrée.''', D, L)

func_block(6, 'LISTE COMPLÈTE DES PERSONNES AUTORISÉES',
'''Au-delà des vérifications à la sortie, le gardien peut consulter à tout moment la liste complète des personnes autorisées pour tous les élèves de l\'école. Il peut rechercher par nom d\'élève pour retrouver rapidement les personnes autorisées pour un enfant spécifique. Cette fonctionnalité lui permet de se préparer, par exemple, lorsqu\'un parent lui a signalé qu\'une personne viendrait chercher son enfant.''', D, L)

func_block(7, 'PROFIL ET ACTIVITÉ RÉCENTE',
'''Le gardien dispose d\'une section profil où il peut consulter son historique d\'activité récente (les scans effectués ces derniers jours) et gérer ses informations personnelles : mise à jour de son numéro de téléphone et changement de son code PIN. Ses informations de connexion sont strictement personnelles et ne doivent jamais être partagées.''', D, L)

recap_box([
    'Tableau de bord de sécurité (statut lockdown, compteurs entrées/sorties)',
    'Scanner d\'entrée QR Code (résultat immédiat : vert / orange / rouge)',
    'Scanner de sortie avec affichage des personnes autorisées + photos',
    'Zoom plein écran sur la photo d\'une personne autorisée',
    'Journal chronologique des scans de la journée',
    'Liste complète des personnes autorisées par élève',
    'Profil et historique d\'activité',
], D, L)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════
h1('Conclusion')
body('''SchoolSafe v3.0 propose un total de 58 fonctionnalités réparties sur ses six profils utilisateurs. Chaque profil est conçu pour correspondre exactement aux besoins et aux responsabilités de son utilisateur, sans surcharger l\'interface avec des fonctionnalités superflues. L\'enseignant ne voit que ce qui concerne sa classe. Le gardien ne voit que ce qui concerne la sécurité. Le parent ne voit que ce qui concerne son enfant.''')
body('''Cette organisation garantit que chaque acteur de la communauté scolaire du Complexe Scolaire Le Sage peut utiliser SchoolSafe efficacement dès les premiers jours, sans formation longue, et qu\'il trouve immédiatement les outils dont il a besoin pour accomplir sa mission.''')

# Tableau récapitulatif final
tbl = doc.add_table(rows=7, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for j, h in enumerate(['Profil', 'Rôle principal', 'Nb de fonctionnalités']):
    c = tbl.rows[0].cells[j]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
    c.paragraphs[0].runs[0].font.color.rgb = BLANC
    set_bg(c, '0A2A5C')

for i, (profil, role, nb, bg) in enumerate([
    ('👨‍🏫 Enseignant',        'Production de l\'information pédagogique',  '11', 'D6EAF8'),
    ('🏫 Direction 1',         'Administration générale et gouvernance',     '14', 'D5F5E3'),
    ('📚 Direction 2',         'Supervision pédagogique et validation',      '9',  'F5E6FF'),
    ('💰 Caisse (Direction 3)', 'Gestion financière',                        '7',  'FDEBD0'),
    ('👨‍👩‍👧 Parent',             'Suivi de la scolarité de l\'enfant',          '10', 'FEF9E7'),
    ('🛡️ Gardien',             'Contrôle des entrées et sorties',            '7',  'FDEDEC'),
]):
    row = tbl.rows[i+1]
    row.cells[0].text = profil; row.cells[1].text = role; row.cells[2].text = nb
    for j in range(3):
        row.cells[j].paragraphs[0].runs[0].font.size = Pt(10)
        set_bg(row.cells[j], bg)
    row.cells[0].paragraphs[0].runs[0].bold = True

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('« Un enfant protégé · Un parent informé »')
r.bold = True; r.italic = True; r.font.size = Pt(14); r.font.color.rgb = BLEU

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SchoolSafe v3.0  ·  Complexe Scolaire Le Sage  ·  Kinshasa, DRC  ·  ' + str(datetime.date.today().year))
r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

out = '/home/user/zalavrai/SchoolSafe_Guide_Profils.docx'
doc.save(out)
print('OK :', out)
